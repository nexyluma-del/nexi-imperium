#!/usr/bin/env python3
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(r"C:\AI\projects\09-video-analyse")
DESKTOP_KI = Path(r"C:\Users\nexil\Desktop\KI")
STATUS_JSON = PROJECT_DIR / "dashboard" / "status.json"
OUT = Path("FINAL-DOKUMENTATION-NEXI-IMPERIUM.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 40)
MUTED = RGBColor(94, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
LINE = "C8D1DC"


def load_status() -> dict[str, Any]:
    if STATUS_JSON.exists():
        return json.loads(STATUS_JSON.read_text(encoding="utf-8"))
    return {}


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color or INK, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_geometry(table, widths: list[int], header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            width = widths[min(col_index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if header and row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
            else:
                shade_cell(cell, WHITE)
            set_cell_margins(cell)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("NEXI IMPERIUM")
    set_run_font(run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Final-Dokumentation")
    set_run_font(run, size=30, color=INK, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(20)
    run = sub.add_run("Lokales KI-System, Video-Pipeline, Memory-KI, RAG, Backup und Wiederaufbau")
    set_run_font(run, size=13.5, color=MUTED)

    meta = [
        ("Stand", datetime.now().strftime("%d.%m.%Y %H:%M")),
        ("Status", "Phase 0 + Phase 1 abgeschlossen, Phase 2 Kernpipeline produktiv nutzbar"),
        ("Default-Wissensmodell", "OpenWebUI: nexi-rag-qwen3-30b / qwen3:30b"),
        ("Dashboard", "http://127.0.0.1:8765"),
        ("Privates Repo", "C:\\AI\\imperium-config"),
    ]
    add_key_value_table(doc, meta, [2200, 7160], header=None)

    add_note(
        doc,
        "Betriebsregel",
        "Keys, Passwoerter, Recovery-Keys und API-Tokens werden nicht in dieser Dokumentation gespeichert. "
        "Sie bleiben in Standard Notes beziehungsweise in der lokalen .env.",
    )


def add_note(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], header=False)
    shade_cell(table.cell(0, 0), LIGHT_GRAY)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: list[int], header: tuple[str, str] | None = None) -> None:
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    if header:
        set_cell_text(table.cell(0, 0), header[0], bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(0, 1), header[1], bold=True, color=DARK_BLUE)
        start = 1
    else:
        start = 0
    for index, (label, value) in enumerate(rows, start=start):
        set_cell_text(table.cell(index, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(index, 1), value)
    set_table_geometry(table, widths, header=bool(header))
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for col, header in enumerate(headers):
        set_cell_text(table.cell(0, col), header, bold=True, color=DARK_BLUE)
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), value)
    set_table_geometry(table, widths, header=True)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_status_overview(doc: Document, status: dict[str, Any]) -> None:
    doc.add_heading("1. Systemstatus", level=1)
    qdrant = status.get("qdrant", {}).get("counts", {})
    sof = status.get("sofinello", {})
    rows = [
        ["Dashboard", "aktiv", "http://127.0.0.1:8765"],
        ["OpenWebUI", "aktiv", "Default-RAG-Modell nexi-rag-qwen3-30b"],
        ["Qdrant video_knowledge", str(qdrant.get("video_knowledge", {}).get("points", "n/a")), "allgemeines Video-Wissen"],
        ["Qdrant open-webui_knowledge", str(qdrant.get("open-webui_knowledge", {}).get("points", "n/a")), "OpenWebUI-RAG-Spiegel"],
        ["Qdrant sofinello_knowledge", str(qdrant.get("sofinello_knowledge", {}).get("points", "n/a")), "Sofinello-Spezialwissen"],
        ["Sofinello Batch B", f"{sof.get('processed', 0)} / {sof.get('total', 0)}", f"Kosten bisher: {sof.get('cost_usd', 0)} USD"],
        ["Failed Videos", str(status.get("failed_videos", {}).get("count", 0)), "laufende Fehlerliste"],
    ]
    add_matrix(doc, ["Bereich", "Status/Wert", "Hinweis"], rows, [2600, 2200, 4560])


def add_architecture(doc: Document) -> None:
    doc.add_heading("2. Architektur", level=1)
    add_key_value_table(
        doc,
        [
            ("Laptop", "HP Omen Max, RTX 5090 Laptop GPU, 64 GB RAM"),
            ("Lokale Runtime", "WSL2 Ubuntu 24.04, Docker Desktop, PowerShell"),
            ("LLM lokal", "Ollama mit qwen3:30b und qwen3:4b"),
            ("Embedding", "nomic-embed-text"),
            ("Vektor-DB", "Qdrant mit video_knowledge, open-webui_knowledge, sofinello_knowledge, memory_voice"),
            ("UI", "OpenWebUI als Chat-Schicht, n8n als Workflow-Schicht"),
            ("Automatisierung", "Telegram-Bot, Scheduler, Python-Skripte"),
            ("Backup", "Restic nach D:\\Restic-Backup plus GitHub-Repo C:\\AI\\imperium-config"),
        ],
        [2300, 7060],
        header=("Komponente", "Beschreibung"),
    )
    add_note(
        doc,
        "Hardware-Doktrin",
        "Phase 1 bleibt lokal auf dem Laptop. Nicht-GPU-Agenten koennen spaeter auf einen Mini-VPS. "
        "GPU-Server werden erst mit stabilem Cashflow sinnvoll.",
    )


def add_paths_and_urls(doc: Document) -> None:
    doc.add_heading("3. Zentrale URLs und Pfade", level=1)
    add_key_value_table(
        doc,
        [
            ("Status-Dashboard", "http://127.0.0.1:8765"),
            ("OpenWebUI", "http://127.0.0.1:3000"),
            ("n8n", "http://127.0.0.1:5678"),
            ("Qdrant", "http://127.0.0.1:6333"),
            ("Ollama", "http://127.0.0.1:11434"),
            ("Projekt", "C:\\AI\\projects\\09-video-analyse"),
            ("GitHub-Repo lokal", "C:\\AI\\imperium-config"),
            ("KI-Ordner", "C:\\Users\\nexil\\Desktop\\KI"),
            ("Instagram Liste", "C:\\Users\\nexil\\Desktop\\Instagram Liste"),
            ("Instagram Videos", "C:\\Users\\nexil\\Desktop\\Instagram Videos"),
            ("Obsidian", "C:\\Users\\nexil\\Documents\\Obsidian-Imperium"),
        ],
        [2500, 6860],
        header=("Name", "Wert"),
    )


def add_pipelines(doc: Document) -> None:
    doc.add_heading("4. Pipelines", level=1)
    add_matrix(
        doc,
        ["Pipeline", "Zweck", "Status"],
        [
            ["URL-Video", "yt-dlp -> Whisper -> Gemini -> Markdown -> Qdrant", "produktiv getestet"],
            ["Batch-Splitter", "Rohlisten nach Themen trennen und Fragen je Video erhalten", "produktiv getestet"],
            ["Bilder/Cross-Check", "Instagram-Posts mit Gemini, Claude und ChatGPT pruefen", "produktiv getestet"],
            ["Lokale Videos", "Ordner Desktop\\Instagram Videos analysieren", "normale Kategorien aktiv"],
            ["Sofinello", "Frame-Upscaling, Gemini, Compliance-Agent, Qdrant", "197/722 verarbeitet, Resume-Guard bereit"],
            ["Telegram Share", "Instagram-Link vom Handy -> Bot -> Analyse -> Rueckmeldung", "aktiv"],
        ],
        [2100, 5100, 2160],
    )
    add_note(
        doc,
        "Sofinello-Regel",
        "Sofinello ist D2-freigegeben, aber jeder oeffentlich bestimmte Output braucht Compliance-Check, "
        "keine Heilversprechen und keine Diagnosen.",
    )


def add_memory_and_sync(doc: Document) -> None:
    doc.add_heading("5. Memory-KI und KI-Sync", level=1)
    add_bullets(
        doc,
        [
            "Memory-KI nutzt Qdrant-Wissen und lokale Projektdateien fuer Rueckfragen, Briefings und Ideenverknuepfungen.",
            "KI-PUSH wird woechentlich gescannt, damit Best-Practices und neue KI-Tutorials in die Strategie einfließen.",
            "KI-Sync-Bridge exportiert Master- oder Topic-Kontexte fuer Claude, ChatGPT und Gemini.",
            "Telegram-Kommandos /memory, /briefing, /links, /sync und /sync-tg sind die mobile Bedienungsschicht.",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("Master-Kontext", "C:\\Users\\nexil\\Desktop\\KI\\sync"),
            ("Memory-Skripte", "memory_query.py, memory_briefing.py, memory_ki_push_scan.py"),
            ("Sync-Skripte", "export_master_context.py, export_topic_context.py"),
            ("Voice-Capture", "voice_capture.py plus Scheduler fuer Review"),
        ],
        [2500, 6860],
        header=("Baustein", "Datei/Pfad"),
    )


def add_operations(doc: Document) -> None:
    doc.add_heading("6. Betrieb im Alltag", level=1)
    add_numbered(
        doc,
        [
            "Dashboard oeffnen und pruefen, ob OpenWebUI, Qdrant, Ollama, n8n und Telegram gruen sind.",
            "Neue Video-Rohlisten in Desktop\\Instagram Liste schreiben und per Splitter trennen lassen.",
            "Batch mit Kostenlimit starten und failed-videos.md beobachten.",
            "Wichtige Erkenntnisse ueber OpenWebUI, Telegram oder Memory-KI abfragen.",
            "Nach groesseren Aenderungen GitHub-Push und Restic-Backup durchfuehren.",
        ],
    )
    add_matrix(
        doc,
        ["Aktion", "Befehl/Ort", "Hinweis"],
        [
            ["Dashboard", "http://127.0.0.1:8765", "alle 60 Sekunden aktualisiert"],
            ["Bot starten", "Desktop\\KI\\scripts\\start-telegram-bot.ps1", "Share-Modus vom Handy"],
            ["Backup", "Desktop\\KI\\scripts\\run-backup.ps1", "Restic-Passwort aus Standard Notes"],
            ["OpenWebUI RAG Sync", "sync_openwebui_qdrant.py", "nach Wissensimport ausfuehren"],
        ],
        [2100, 4200, 3060],
    )


def add_backup_rebuild(doc: Document) -> None:
    doc.add_heading("7. Backup und Wiederaufbau", level=1)
    add_bullets(
        doc,
        [
            "Restic sichert Desktop, Documents, Pictures, Videos und C:\\AI; grosse Modelle sind ausgeschlossen.",
            "Docker-Volumes werden ueber Export-Skripte separat sicherbar gemacht.",
            "GitHub-Repo enthaelt reproduzierbare Skripte, Policies und Dokumentation.",
            "BitLocker-Recovery-Key und Restic-Passwort sind in Standard Notes gespeichert.",
        ],
    )
    add_matrix(
        doc,
        ["Rebuild-Schritt", "Quelle", "Ziel"],
        [
            ["Basis", "Windows + WSL2 + Docker Desktop", "lokale Laufzeit"],
            ["Repo", "nexyluma-del/nexi-imperium", "C:\\AI\\imperium-config"],
            ["Projekt", "Restic oder Arbeitskopie", "C:\\AI\\projects\\09-video-analyse"],
            ["Wissen", "Qdrant Snapshot / Restic", "Qdrant Collections"],
            ["UI", "OpenWebUI Volume", "open-webui"],
            ["Workflows", "n8n Export", "n8n"],
        ],
        [2300, 3600, 3460],
    )


def add_costs_and_guardrails(doc: Document) -> None:
    doc.add_heading("8. Kosten, Datenklassen und Guardrails", level=1)
    add_key_value_table(
        doc,
        [
            ("D0/D1", "lokal oder Cloud unkritisch"),
            ("D2", "Cloud erlaubt, wenn oeffentliche Inhalte und Freigabe vorhanden"),
            ("D3/D4", "keine Cloud ohne explizite Freigabe"),
            ("Batch-Kosten", "immer mit Max-Budget starten"),
            ("Sofinello", "Compliance-Agent Pflicht; keine Heilversprechen, keine Diagnosen"),
            ("Server-Skalierung", "erst VPS bei Web-Cashflow, GPU-Server erst bei stabilem Einkommen"),
        ],
        [2300, 7060],
        header=("Regel", "Bedeutung"),
    )


def add_backlog(doc: Document) -> None:
    doc.add_heading("9. Backlog und offene Beobachtung", level=1)
    add_bullets(
        doc,
        [
            "Aufgabe 020 Chief-Workflow-Template bleibt Backlog.",
            "Sofinello-Batch B laeuft per Resume-Guard weiter, sobald Quota wieder greift.",
            "C-Modus Full-Video-Upscaling bleibt fuer 10-20 ausgewaehlte Verpackungs-/OCR-Clips reserviert.",
            "Memory-KI kann spaeter staerker mit Obsidian, KI-PUSH und Chief-Agenten verzahnt werden.",
            "Bei Cashflow: Mini-VPS fuer Nicht-GPU-Agenten, spaeter GPU-Server.",
        ],
    )


def add_appendix(doc: Document) -> None:
    doc.add_heading("10. Wichtigste Befehle", level=1)
    commands = [
        ("Dienste", "docker ps"),
        ("Dashboard", 'powershell.exe -NoProfile -ExecutionPolicy Bypass -File "C:\\Users\\nexil\\Desktop\\KI\\scripts\\start-status-dashboard.ps1"'),
        ("OpenWebUI-RAG-Test", 'wsl.exe -d Ubuntu-24.04 -- bash -lc "cd /mnt/c/AI/projects/09-video-analyse && .venv/bin/python scripts/test_openwebui_default_rag.py"'),
        ("Telegram Status", 'wsl.exe -d Ubuntu-24.04 -- bash -lc "cd /mnt/c/AI/projects/09-video-analyse && .venv/bin/python scripts/telegram_status.py"'),
        ("Memory Frage", 'wsl.exe -d Ubuntu-24.04 -- bash -lc "cd /mnt/c/AI/projects/09-video-analyse && .venv/bin/python scripts/memory_query.py --fast \\"Was ist der naechste Schritt?\\""'),
        ("Git Push", "cd C:\\AI\\imperium-config && git status && git add . && git commit -m \"Update\" && git push"),
    ]
    add_key_value_table(doc, commands, [2200, 7160], header=("Aktion", "Befehl"))


def set_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = ""
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = header.add_run("Nexi Imperium - Final-Dokumentation")
        set_run_font(r, size=9, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = footer.add_run("Privates Betriebsdokument - keine Secrets enthalten")
        set_run_font(r2, size=8.5, color=MUTED)


def build() -> Path:
    status = load_status()
    doc = Document()
    apply_styles(doc)
    set_header_footer(doc)
    add_cover(doc)
    doc.add_page_break()
    add_status_overview(doc, status)
    add_architecture(doc)
    add_paths_and_urls(doc)
    doc.add_page_break()
    add_pipelines(doc)
    add_memory_and_sync(doc)
    add_operations(doc)
    doc.add_page_break()
    add_backup_rebuild(doc)
    add_costs_and_guardrails(doc)
    add_backlog(doc)
    doc.add_page_break()
    add_appendix(doc)
    doc.core_properties.title = "Nexi Imperium Final-Dokumentation"
    doc.core_properties.subject = "Lokales KI-System, RAG, Pipelines, Backup und Wiederaufbau"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path.resolve())
